import flet as ft
import markdown2
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
import os
import re

# Функция для конвертации Markdown в HTML
def markdown_to_html(markdown_text):
    return markdown2.markdown(markdown_text, extras=["tables"])

# Функция для удаления недопустимых символов из имени файла
def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

# Функция для создания уникального имени файла
def get_unique_filename(directory, base_name, extension=".docx"):
    if not base_name:
        base_name = "Converted Document"
    sanitized_name = sanitize_filename(base_name)
    full_path = os.path.join(directory, f"{sanitized_name}{extension}")
    if not os.path.exists(full_path):
        return full_path
    copy_number = 1
    while True:
        new_name = f"{sanitized_name} (копия {copy_number})"
        new_path = os.path.join(directory, f"{new_name}{extension}")
        if not os.path.exists(new_path):
            return new_path
        copy_number += 1

# Функция для конвертации HTML в Word с поддержкой текста и нескольких таблиц
def html_to_word(html, output_file="output.docx"):
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Проверяем наличие стиля 'BlackHeading' перед добавлением
    try:
        doc.styles['BlackHeading']
    except KeyError:
        style = doc.styles.add_style('BlackHeading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

    for element in soup.contents:
        if element.name in ["h1", "h2", "h3"]:
            doc.add_paragraph(element.get_text(), style='BlackHeading')

        elif element.name == "table":
            rows = element.find_all("tr")
            word_table = doc.add_table(rows=1, cols=len(rows[0].find_all(["th", "td"])))

            # Заполняем заголовки таблицы
            hdr_cells = word_table.rows[0].cells
            for i, header in enumerate(rows[0].find_all(["th", "td"])):
                hdr_cells[i].text = header.get_text()

            # Заполняем строки таблицы
            for row in rows[1:]:
                cells = row.find_all("td")
                word_row = word_table.add_row().cells
                for i, cell in enumerate(cells):
                    word_row[i].text = cell.get_text()

            # Устанавливаем границы для всех ячеек таблицы
            for row in word_table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    borders = parse_xml(
                        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        r'<w:top w:val="single" w:sz="4"/>'
                        r'<w:left w:val="single" w:sz="4"/>'
                        r'<w:bottom w:val="single" w:sz="4"/>'
                        r'<w:right w:val="single" w:sz="4"/>'
                        r'</w:tcBorders>'
                    )
                    tcPr.append(borders)

        elif element.name == "p" or element.name is None:
            text = element.get_text().strip() if element.name == "p" else element.strip()
            if text:
                doc.add_paragraph(text)

    doc.save(output_file)
    return True, f"Документ успешно сохранён как '{output_file}'."


# Основная функция для интерфейса Flet
def main(page: ft.Page):
    page.title = "Конвертация Markdown в Word"
    page.scroll = "adaptive"

    markdown_input = ft.TextField(
        label="Введите текст в формате Markdown",
        multiline=True,
        min_lines=10,
        max_lines=20,
        expand=True
    )

    result_message = ft.Text()

    file_name_input = ft.TextField(label="Имя файла:", expand=True)
    file_path_input = ft.TextField(label="Путь сохранения:", expand=True)

    def convert_to_word(e):
        markdown_text = markdown_input.value
        if not markdown_text.strip():
            result_message.value = "Введите Markdown текст."
            page.update()
            return

        html = markdown_to_html(markdown_text)
        save_path = file_path_input.value.strip() or os.getcwd()
        file_name = get_unique_filename(save_path, file_name_input.value)
        success, message = html_to_word(html, file_name)

        result_message.value = message
        page.update()

    convert_button = ft.ElevatedButton("Конвертировать", on_click=convert_to_word)

    page.add(markdown_input, file_name_input, file_path_input, convert_button, result_message)

# Запуск приложения Flet
ft.app(target=main)
